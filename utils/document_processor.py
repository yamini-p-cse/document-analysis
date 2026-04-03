import PyPDF2
from docx import Document
from typing import Optional
import os
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    @staticmethod
    def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
        """Extract text from PDF with better error handling"""
        if not os.path.exists(pdf_path):
            logger.error(f"PDF file not found: {pdf_path}")
            return None
            
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if not pdf_reader.pages:
                    logger.warning("PDF has no pages")
                    return None
                    
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text() or ""
                    text += f"\n--- Page {page_num} ---\n{page_text}\n\n"
                    
        except Exception as e:
            logger.error(f"PDF Error [{pdf_path}]: {str(e)}")
            return None
            
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(docx_path: str) -> Optional[str]:
        """Extract text from DOCX with better formatting"""
        if not os.path.exists(docx_path):
            logger.error(f"DOCX file not found: {docx_path}")
            return None
            
        text = ""
        try:
            doc = Document(docx_path)
            
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
            
            # Tables
            for table_num, table in enumerate(doc.tables, 1):
                text += f"\n--- Table {table_num} ---\n"
                for row_idx, row in enumerate(table.rows):
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += f"Row {row_idx + 1}: {row_text}\n"
                text += "\n"
                
        except Exception as e:
            logger.error(f"DOCX Error [{docx_path}]: {str(e)}")
            return None
            
        return text.strip()
    
    @staticmethod
    def extract_text(file_path: str) -> Optional[str]:
        """Extract text from ANY supported format"""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
            
        ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"Processing {ext}: {os.path.basename(file_path)}")
        
        if ext == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif ext in ['.txt', '.text']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                # Try latin-1 fallback
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read().strip()
            except Exception as e:
                logger.error(f"TXT Error: {str(e)}")
                return None
        elif ext.lower() in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            logger.warning("OCR not implemented yet - install easyocr/pytesseract for image support")
            return None
        
        logger.warning(f"Unsupported format: {ext}")
        return None

# Usage Example
if __name__ == "__main__":
    processor = DocumentProcessor()
    text = processor.extract_text("sample.pdf")
    print(text[:500] if text else "No text extracted")